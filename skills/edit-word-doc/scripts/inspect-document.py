# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx>=1.1"]
# ///
"""Dump the structure of a .docx so edits can be planned safely.

Reports sections, styles in use, paragraph/table/image counts, headers and
footers, and whether tracked changes are present (python-docx cannot edit
those), so the agent knows the document's real shape before touching it.
"""

import sys
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


def main() -> int:
    if len(sys.argv) != 2:
        print(f"Usage: uv run {sys.argv[0]} <document.docx>")
        return 2
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"Error: file not found: {path}")
        return 1
    if path.suffix.lower() != ".docx":
        print(f"Error: {path.suffix} is not .docx; legacy .doc needs conversion first.")
        return 1

    doc = Document(str(path))
    paragraphs = doc.paragraphs
    print(f"Document: {path}")
    print(f"Sections: {len(doc.sections)}, paragraphs: {len(paragraphs)}, tables: {len(doc.tables)}, inline images: {len(doc.inline_shapes)}")

    style_counts = Counter(p.style.name for p in paragraphs if p.text.strip())
    print(f"Paragraph styles in use: {dict(style_counts.most_common(10))}")

    headings = [(p.style.name, p.text.strip()[:60]) for p in paragraphs if p.style.name.startswith("Heading") and p.text.strip()]
    if headings:
        print("Outline:")
        for style, text in headings[:25]:
            indent = "  " * (int(style.split()[-1]) if style.split()[-1].isdigit() else 1)
            print(f"  {indent}{text}")

    for i, section in enumerate(doc.sections):
        hdr = " | ".join(p.text.strip() for p in section.header.paragraphs if p.text.strip())
        ftr = " | ".join(p.text.strip() for p in section.footer.paragraphs if p.text.strip())
        if hdr or ftr:
            print(f"Section {i}: header={hdr!r} footer={ftr!r}")

    # python-docx silently ignores tracked changes; detect them from the raw XML
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml")
    if b"<w:ins " in xml or b"<w:del " in xml:
        print("WARNING: document contains tracked changes (w:ins/w:del). python-docx does not model them; paragraph text may be incomplete and edits will not touch the revisions. Resolve them in Word first, or manipulate the XML deliberately.")
    if b"<w:comment" in xml or "word/comments.xml" in zf.namelist():
        print("NOTE: document contains comments.")

    available = [s.name for s in doc.styles if s.type is not None][:30]
    print(f"Styles available in template: {available}{' ...' if len(list(doc.styles)) > 30 else ''}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
